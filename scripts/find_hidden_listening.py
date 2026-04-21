"""Find listening transcriptions for lessons that lack a `Transcription` header.

Pattern discovered in the docx: for many lessons, the listening section lives
inside `texte B` and is preceded by `Écoutez l'enregistrement` / `Écoutez le
document` exercise prompts, with the actual transcription appearing after the
Corrigé answers and before the next `知识点` vocab block. Sometimes a
`Découvrez` label marks the start explicitly.

This script targets the 17 lessons audited as `reading only` (no listening
found). For each lesson:
  1. Find texte B bounds
  2. Within texte B, find Écoutez-triggered zones
  3. In each zone, collect long French prose paragraphs that aren't exercise
     prompts / Corrigé answers / vocab
  4. Dedupe against current books.ts cards
  5. Write new l* cards with a s_listen marker

Prints dry-run summary unless --apply is passed.
"""
from __future__ import annotations
import argparse
import io
import re
import sys
import urllib.request
import json
from pathlib import Path

import docx

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

ROOT = Path(r"C:/Users/racwang/hoopy/french-flashcard")
DOCX = Path(r"C:/Users/racwang/Downloads/B1讲义完整版.pdf+改25.5.31775191233.docx")
BOOKS = ROOT / "src" / "data" / "books.ts"

# Lessons classified as reading-only but believed to have listening somewhere
# Pulled from the earlier audit
TARGETS = [
    "u1l3", "u2l5", "u2l7", "u3l9", "u3l10",
    "u6l21", "u7l26", "u7l27",
    "u9l33", "u9l34",
    "u10l38", "u10l39",
    "u11l41", "u11l42", "u11l43",
    "u12l46", "u12l47",
]

# Lesson bounds — same scheme as earlier scripts, discovered via docx heading scan
# I'll derive these from the docx by finding first-heading positions per lesson
LESSON_RE = re.compile(r"^(U\d+L\d+)\b")


def lesson_bounds(paragraphs):
    firsts = {}
    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        m = LESSON_RE.match(t)
        if m and m.group(1).lower() not in firsts:
            firsts[m.group(1).lower()] = i
    order = sorted(
        firsts.keys(),
        key=lambda x: (int(re.search(r"u(\d+)", x).group(1)), int(re.search(r"l(\d+)", x).group(1))),
    )
    bounds = {}
    for k, lid in enumerate(order):
        s = firsts[lid]
        e = firsts[order[k + 1]] if k + 1 < len(order) else len(paragraphs)
        bounds[lid] = (s, e)
    return bounds


def is_french_prose(t: str) -> bool:
    if len(t) < 40:
        return False
    latin = sum(1 for c in t if c.isalpha() and ord(c) < 0x2E80)
    cjk = sum(1 for c in t if 0x4E00 <= ord(c) <= 0x9FFF)
    if cjk > latin // 5:
        return False
    return latin > 30


SKIP_START = re.compile(
    r"^("
    r"知识点|Corrigé|Exercice|Vocabulaire|Prononcez|Devoir|grammaire|讲解|课文讲解|"
    r"Observez|Lisez|Dites|Complétez|Répondez|Imaginez|Faites|Présentez|Écoutez|Ecoutez|"
    r"Mettez|Associez|Trouvez|Choisissez|Relevez|Expliquez|Racontez|Classez|Cochez|Indiquez|"
    r"Discutez|Interviewez|Décrivez|Résumez|Comparez|Préparez|Commentez|Cherchez|"
    r"Repérez|Notez|Identifiez|Formulez|Reformulez|Transformez|Citez|Analysez|Nommez|Recherchez|"
    r"Rédigez|Composez|Retrouvez|Soulignez|Donnez|Employez|Utilisez|Créez|Jouez|Écrivez|"
    r"Distinguez|Poursuivez|Continuez|Ex\s*:|"
    r"structure|一、|二、|三、|四、|五、|六、|预习|自主|Transcription|听力|"
    r"Vrai|Faux|Document\b|document\b|Découvrez|"
    r"[a-hA-H][\s.)）、]|[0-9]+[.．、]|\(\d+\)|【|®"
    r")",
    re.I,
)


def find_texte_b_bounds(paragraphs, lesson_start: int, lesson_end: int) -> list[tuple[int, int]]:
    """Find all texte B sub-ranges inside a lesson."""
    ranges = []
    current_start = None
    for i in range(lesson_start, lesson_end):
        t = paragraphs[i].text.strip()
        if not t:
            continue
        # Match "U?L? ... - texte B" or "texte B" or "– texte B"
        if re.search(r"texte\s*B\b", t, re.I):
            if current_start is not None:
                ranges.append((current_start, i))
            current_start = i
        elif re.search(r"texte\s*A\b|grammaire|texte\s*\(?1\)?|texte\s*\(?2\)?", t, re.I):
            if current_start is not None:
                ranges.append((current_start, i))
                current_start = None
    if current_start is not None:
        ranges.append((current_start, lesson_end))
    return ranges


def extract_transcription(paragraphs, s: int, e: int) -> list[str]:
    """Within a texte B range, find long French prose that isn't exercise or vocab."""
    kept = []
    in_vocab = False
    for i in range(s, e):
        t = paragraphs[i].text.strip()
        if not t:
            continue
        if t.startswith(("知识点", "讲解")):
            in_vocab = True
            continue
        if re.match(r"^U\d+L\d+\b|^texte\b", t, re.I):
            in_vocab = False
            continue
        if in_vocab:
            continue
        if SKIP_START.match(t):
            continue
        if is_french_prose(t):
            kept.append(t)
    return kept


def split_sentences(paragraphs: list[str]) -> list[str]:
    text = " ".join(p.strip() for p in paragraphs if p.strip())
    text = re.sub(r"\s+", " ", text)
    for abbr in ["M.", "Mme.", "Mlle.", "Dr.", "St.", "Ste.", "pp.", "p.", "etc.", "ex.", "cf.", "i.e.", "e.g.", "vs."]:
        text = text.replace(abbr, abbr.replace(".", "\x00"))
    sents = re.split(r"(?<=[.!?…])\s+(?=[A-ZÀ-ÿ«\"'“‘\(])", text)
    sents = [s.replace("\x00", ".").strip() for s in sents if s.strip()]
    return [s for s in sents if len(s) >= 10]


def norm(s: str) -> str:
    return re.sub(r"[^\w]+", "", s.lower())


def fetch_current_fronts():
    d = json.loads(urllib.request.urlopen("http://localhost:3001/api/editor").read())
    result = {}
    for l in d[0]["lessons"]:
        result[l["id"]] = {norm(c["front"])[:30] for c in l["cards"] if not c["id"].startswith("s_")}
    return result


CARD_RE = re.compile(
    r"\{\s*id:\s*`([^`]+)`\s*,\s*front:\s*`([^`]*)`\s*,\s*back:\s*`([^`]*)`\s*\}",
)
LESSON_BLOCK_RE = re.compile(
    r"(\{\s*id:\s*`(u\d+l\d+)`\s*,\s*title:\s*`[^`]+`\s*,\s*cards:\s*\[)(.*?)(\s*\]\s*,?\s*\})",
    re.S,
)


def esc(s: str) -> str:
    return s.replace("\\", "\\\\").replace("`", "\\`").replace("${", "\\${")


def card_literal(cid: str, front: str, back: str) -> str:
    return f"{{ id: `{esc(cid)}`, front: `{esc(front)}`, back: `{esc(back)}` }}"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--apply", action="store_true")
    parser.add_argument("--show", type=str, default=None, help="Dump one lesson's candidates fully")
    args = parser.parse_args()

    d = docx.Document(str(DOCX))
    bounds = lesson_bounds(d.paragraphs)
    existing = fetch_current_fronts()

    per_lesson_sentences: dict[str, list[str]] = {}
    per_lesson_report: list[tuple[str, int, int, int]] = []

    for lid in TARGETS:
        if lid not in bounds:
            per_lesson_report.append((lid, 0, 0, 0))
            continue
        s, e = bounds[lid]
        tb_ranges = find_texte_b_bounds(d.paragraphs, s, e)
        raw = []
        for rs, re_ in tb_ranges:
            raw.extend(extract_transcription(d.paragraphs, rs, re_))
        sents = split_sentences(raw)
        # Dedupe against current books.ts
        bk = existing.get(lid, set())
        new = []
        for st in sents:
            k = norm(st)[:30]
            if len(k) < 10 or k in bk:
                continue
            new.append(st)
            bk.add(k)
        per_lesson_sentences[lid] = new
        per_lesson_report.append((lid, len(raw), len(sents), len(new)))

    print(f"{'lesson':<8} {'raw-para':>8} {'sents':>6} {'new':>5}")
    print("-" * 32)
    for lid, rp, s, n in per_lesson_report:
        print(f"{lid:<8} {rp:>8} {s:>6} {n:>5}")

    if args.show:
        lid = args.show
        print(f"\n=== {lid} candidates ===")
        for i, s in enumerate(per_lesson_sentences.get(lid, [])):
            print(f"  l{i+1}: {s[:140]}")

    if not args.apply:
        print("\n(dry-run — use --apply to write)")
        return

    # Apply: for each lesson, append s_listen marker + new l* cards to its block
    content = BOOKS.read_text(encoding="utf-8")

    def rebuild(m):
        head, lid, body, tail = m.group(1), m.group(2), m.group(3), m.group(4)
        if lid not in per_lesson_sentences or not per_lesson_sentences[lid]:
            return m.group(0)

        new_sentences = per_lesson_sentences[lid]
        existing_cards = [(cm.group(1), cm.group(2), cm.group(3)) for cm in CARD_RE.finditer(body)]

        # Check if s_listen already exists in this lesson
        has_listen = any(c[0] == "s_listen" for c in existing_cards)

        # Build new tail: (maybe s_listen) + l* cards
        new_cards = []
        if not has_listen:
            new_cards.append(("s_listen", "【Compréhension orale · 听力】", ""))
        for i, s in enumerate(new_sentences, 1):
            new_cards.append((f"l{i}", s, ""))

        # Append new_cards to existing_cards
        all_cards = existing_cards + new_cards

        # Render
        lines = []
        for cid, front, back in all_cards:
            lines.append("          " + card_literal(cid, front, back) + ",")
        new_body = "\n" + "\n".join(lines) + "\n        "
        return head + new_body + tail

    new_content = LESSON_BLOCK_RE.sub(rebuild, content)
    BOOKS.write_text(new_content, encoding="utf-8")
    total_new = sum(len(v) for v in per_lesson_sentences.values())
    print(f"\nApplied: +{total_new} l* cards across {sum(1 for v in per_lesson_sentences.values() if v)} lessons")


if __name__ == "__main__":
    main()
